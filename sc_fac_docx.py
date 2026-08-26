"""分判最終結算 Word 匯出（供版面微調 · 對照調整 PDF）"""
from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from reportlab.lib.units import mm

from sc_fac_pdf import (
    BODY_EN_FONT_PT,
    BODY_FONT_PT,
    DEFAULT_SC_FAC_THEME,
    HDR_COL_WIDTHS,
    P1_BODY_EN_FONT_PT,
    P1_BODY_FONT_PT,
    _money_contra_amt,
    _money_split,
    _variations_settle_labels,
    _vo_total_label_parts,
    normalize_sc_fac_theme,
)

FONT_NAME = 'Microsoft JhengHei'
COLOR_RED = RGBColor(0xDC, 0x26, 0x26)


def _money_acct(val) -> str:
    try:
        n = float(val or 0)
    except (TypeError, ValueError):
        return '—'
    if n == 0:
        return 'HK$0.00'
    abs_s = f'HK${abs(n):,.2f}'
    return f'({abs_s})' if n < 0 else abs_s


def _set_cell_text(cell, text, *, bold=False, underline=False, strike=False, red=False, align=None, size=BODY_FONT_PT):
    cell.text = ''
    p = cell.paragraphs[0]
    if align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text if text is not None else ''))
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline
    run.font.strike = strike
    if red:
        run.font.color.rgb = COLOR_RED


def _set_cell_vo_total_label(cell, total, *, size=BODY_FONT_PT):
    cell.text = ''
    p = cell.paragraphs[0]
    for text, strike in _vo_total_label_parts(total):
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(size)
        run.bold = True
        run.font.strike = strike


def _cell_border(cell, *, top=False, bottom=False, bottom_double=False):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    if top:
        el = OxmlElement('w:top')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    if bottom or bottom_double:
        el = OxmlElement('w:bottom')
        el.set(qn('w:val'), 'double' if bottom_double else 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tcPr.append(borders)


def _add_para(doc, text, *, center=False, underline=False, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline
    return p


def _landscape_section(doc):
    sec = doc.add_section()
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Mm(15)
    sec.bottom_margin = Mm(28)
    sec.left_margin = Mm(15)
    sec.right_margin = Mm(15)
    return sec


def _portrait_section(doc):
    sec = doc.add_section()
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Mm(15)
    sec.bottom_margin = Mm(28)
    sec.left_margin = Mm(15)
    sec.right_margin = Mm(15)
    return sec


def _appendix_head(doc, h, appendix, section_title):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = True
    _set_cell_text(tbl.rows[0].cells[0], h.get('main_contract_works', '—'), underline=True, size=10)
    _set_cell_text(tbl.rows[0].cells[1], appendix, bold=True, underline=True, align='right', size=11)
    doc.add_paragraph()
    _add_para(doc, section_title, underline=True, bold=True, size=11, space_after=8)


def _appendix_table_vo(doc, items, total):
    hdr = ['No.', 'AI REF', 'QUO. REF.', 'DESCRIPTION', 'QTY', 'UNIT', 'RATE', 'AMOUNT']
    rows = items or []
    extra = 2  # spacer + total
    tbl = doc.add_table(rows=1 + max(len(rows), 1) + extra, cols=8)
    for j, c in enumerate(hdr):
        cell = tbl.rows[0].cells[j]
        align = 'center' if j in (4, 5, 6, 7) else None
        _set_cell_text(cell, c, bold=True, align=align)
        _cell_border(cell, bottom=True)
    data_end = 1 + max(len(rows), 1)
    for i, v in enumerate(rows, 1):
        vals = [
            str(i),
            v.get('ai_ref', ''),
            v.get('quo_ref', ''),
            v.get('description', ''),
            str(v.get('qty', '')),
            v.get('unit', ''),
            _money_acct(v.get('rate')),
            _money_acct(v.get('amount')),
        ]
        for j, txt in enumerate(vals):
            red = str(txt).startswith('(')
            if j in (4, 5, 6):
                align = 'center'
            elif j == 7:
                align = 'right'
            else:
                align = None
            _set_cell_text(tbl.rows[i].cells[j], txt, red=red, align=align)
    if not rows:
        for j in range(8):
            if j in (4, 5, 6):
                align = 'center'
            elif j == 7:
                align = 'right'
            else:
                align = None
            _set_cell_text(tbl.rows[1].cells[j], '—', align=align)
    spacer = data_end
    for j in (5, 6, 7):
        _cell_border(tbl.rows[spacer].cells[j], bottom=True)
    total_row = spacer + 1
    lbl = tbl.rows[total_row].cells[3]
    lbl.merge(tbl.rows[total_row].cells[6])
    _set_cell_vo_total_label(lbl, total)
    amt = _money_acct(total)
    _set_cell_text(
        tbl.rows[total_row].cells[7], amt,
        bold=True, red=amt.startswith('('), align='right',
    )
    _cell_border(tbl.rows[total_row].cells[7], bottom=True, bottom_double=True)
    doc.add_paragraph()


def _appendix_table_contra(doc, items, total, total_label):
    hdr = ['NO', 'DESCRIPTION', '', 'AMOUNT']
    rows = items or []
    extra = 2
    tbl = doc.add_table(rows=1 + max(len(rows), 1) + extra, cols=4)
    for j, c in enumerate(hdr):
        if not c:
            continue
        cell = tbl.rows[0].cells[j]
        _set_cell_text(cell, c, bold=True, align='center' if j == 3 else None)
        _cell_border(cell, bottom=True)
    data_end = 1 + max(len(rows), 1)
    for i, v in enumerate(rows, 1):
        no = v.get('ai_ref') or str(i)
        amt = _money_contra_amt(v.get('amount'))
        _set_cell_text(tbl.rows[i].cells[0], no)
        _set_cell_text(tbl.rows[i].cells[1], v.get('description', ''))
        _set_cell_text(tbl.rows[i].cells[3], amt, align='right')
    if not rows:
        _set_cell_text(tbl.rows[1].cells[0], '—')
        _set_cell_text(tbl.rows[1].cells[1], '—')
        _set_cell_text(tbl.rows[1].cells[3], '—', align='right')
    spacer = data_end
    _cell_border(tbl.rows[spacer].cells[3], bottom=True)
    total_row = spacer + 1
    _set_cell_text(tbl.rows[total_row].cells[1], total_label, bold=True)
    amt = _money_contra_amt(total)
    _set_cell_text(tbl.rows[total_row].cells[3], amt, bold=True, align='right')
    _cell_border(tbl.rows[total_row].cells[3], bottom=True, bottom_double=True)
    doc.add_paragraph()


def _settlement_classic(doc, st):
    paid_en = (
        f'(Less): Total Value Previous Paid\n'
        f'(as at {st.get("paid_as_at_display") or "—"})'
    )
    ob_en = (
        f'Outstanding Balance (as at {st.get("paid_as_at_display") or "—"})'
    )
    vo_zh, vo_en = _variations_settle_labels(st['variations'])
    rows = [
        ('分判合約總額', 'Original Sub-contract Sum:', st['original_sum'], False, False),
        (vo_zh, vo_en, st['variations'], False, False),
        ('結算工程總額', 'Final Sub-Contract Sum:', st['final_sum'], True, True),
        ('(減) 已付工程額', paid_en, -abs(st['total_paid']), False, False),
    ]
    for line in st.get('deduction_lines') or []:
        rows.append((
            f'(減) {line.get("label_zh", "扣款")}',
            f'(Less): {line.get("label_en", "Deduction")}',
            line.get('amount', 0),
            False, False,
        ))
    rows.append(('', '', None, False, False))
    rows.append(('', '', None, False, False))
    rows.append(('剩餘應付工程額', ob_en, st['outstanding'], True, True))
    rows.append(('根據上述餘應付工程額/保固金，按以下期數支付:', '', None, False, False))
    rows.append((
        st.get('final_payment_label_zh', '第一期糧款（尾款）'),
        st.get('final_payment_label_en', 'Payment IP01 (Final Payment)'),
        st.get('final_payment_amount', 0),
        False, False,
    ))
    rows.append((
        '剩餘應付工程額', 'Total Outstanding Payment',
        st.get('total_outstanding', 0), True, True,
    ))

    tbl = doc.add_table(rows=len(rows), cols=4)
    tbl.autofit = True
    for i, (zh, en, amt, bold, ruled) in enumerate(rows):
        _set_cell_text(tbl.rows[i].cells[0], zh, bold=bold, size=P1_BODY_FONT_PT)
        _set_cell_text(tbl.rows[i].cells[1], en or '', bold=bold, size=P1_BODY_EN_FONT_PT)
        if amt is None:
            tbl.rows[i].cells[0].merge(tbl.rows[i].cells[1])
            continue
        cur, num, red = _money_split(amt)
        _set_cell_text(tbl.rows[i].cells[2], cur, bold=bold, size=P1_BODY_FONT_PT)
        _set_cell_text(tbl.rows[i].cells[3], num, bold=bold, red=red, align='right', size=P1_BODY_FONT_PT)
        if ruled:
            _cell_border(tbl.rows[i].cells[3], top=True, bottom=True, bottom_double=True)
    doc.add_paragraph()


def _page1_statement(doc, data, theme):
    h = data['header']
    st = data['settlement']
    sub_lines = [x for x in (h.get('subcontractor_zh'), h.get('subcontractor_en')) if x]
    if not sub_lines:
        sub_lines = [h.get('subcontractor') or '—']
    from sc_contract_ref import display_p1_sub_contract_no
    sub_contract_no = display_p1_sub_contract_no(h.get('sub_contract_no'))

    _add_para(doc, '工程帳目總結算', center=True, bold=True, size=16, space_after=4)
    if theme == 'classic':
        _add_para(doc, 'Statement of Final Accounts', center=True, underline=True, size=11, space_after=12)
    else:
        _add_para(doc, 'Statement of Final Accounts', center=True, size=11, space_after=12)
    doc.add_paragraph()

    fld = doc.add_table(rows=2, cols=4)
    fld.autofit = False
    for i, w in enumerate(HDR_COL_WIDTHS):
        fld.columns[i].width = Mm(w / mm)
    row_data = [
        [
            ('總承判合約/工程編號:\nMain Contract / Job No.:', h.get('main_contract_works', '—')),
            ('分判合約編號:\nSub-Contracts No.:', sub_contract_no),
        ],
        [
            ('分判商:\nSub-contractor:', '\n'.join(sub_lines)),
            ('分判合約:\nSub-Contract Works:', h.get('sc_works', '—')),
        ],
    ]
    for r, pair in enumerate(row_data):
        for c, (lab, val) in enumerate(pair):
            _set_cell_text(fld.rows[r].cells[c * 2], lab, size=P1_BODY_EN_FONT_PT)
            cell = fld.rows[r].cells[c * 2 + 1]
            _set_cell_text(cell, val, size=P1_BODY_FONT_PT)
            if theme == 'classic':
                _cell_border(cell, bottom=True)
    doc.add_paragraph()

    if theme == 'classic':
        _settlement_classic(doc, st)
    else:
        _settlement_classic(doc, st)

    doc.add_paragraph()
    _add_para(doc, '— 簽名區（P1 內部簽署 · 編制者／合約部／項目部／總經理）—', size=8, space_after=4)


def generate_sc_fac_docx(data: dict, theme: str | None = None) -> bytes:
    """P1 直向 · P2 橫向 VO · P3 直向 Contra（classic 完整版面）"""
    theme = normalize_sc_fac_theme(theme)
    h = data.get('header') or {}

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Mm(15)
    sec.bottom_margin = Mm(46 if theme == 'classic' else 28)
    sec.left_margin = Mm(15)
    sec.right_margin = Mm(15)

    _page1_statement(doc, data, theme)

    doc.add_page_break()
    _landscape_section(doc)
    _appendix_head(doc, h, 'APPENDIX I', 'SUMMARY OF VARIATIONS')
    _appendix_table_vo(
        doc,
        data.get('variations') or [],
        data.get('variations_total', 0),
    )
    _add_para(doc,
        '分判商同意上述後加工程結算所詳列之帳目正確無誤。分判商本人/本公司亦承諾不會再向美博工程服務有限公司根據上述後加工程作出任何索償。',
        size=BODY_FONT_PT, space_after=6)

    doc.add_page_break()
    _portrait_section(doc)
    _appendix_head(doc, h, 'APPENDIX II', 'SUMMARY OF CONTRA CHARGE')
    _appendix_table_contra(
        doc,
        data.get('contra_charges') or [],
        data.get('contra_charges_total', 0),
        'Net Contra Charges Carried to Final Account',
    )
    _add_para(doc,
        '分判商同意上述所詳列之帳目正確無誤。分判商本人/本公司亦承諾不會再向美博工程服務有限公司根據上述之支項目作出任何索償。',
        size=BODY_FONT_PT, space_after=6)
    _add_para(doc, '— 頁底簽章區（Mepork · 分判商）—', size=8)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
