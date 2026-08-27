"""QS 地盤財務匯報 Word 匯出（版面對照 QS 交回 Word · 2026-08-14）"""
from __future__ import annotations

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from reportlab.lib.units import mm as rl_mm

from qs_report_pdf import (
    STATUS_LABELS,
    _attention_items,
    _category_rows,
    _company_primary,
    _money,
    _pct,
    _plain,
    _project_titles,
    _sc_category,
)

_FONT = 'Microsoft JhengHei'
_LABEL_FILL = 'F1F5F9'
_HDR_FILL = '334155'
_HDR_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
_TITLE_COLOR = RGBColor(0x1E, 0x3A, 0x5F)
_MUTED_COLOR = RGBColor(0x64, 0x74, 0x8B)


def _run_font(run, *, size=9, bold=False, color=None):
    run.font.name = _FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for tag in ('ascii', 'hAnsi', 'eastAsia'):
        rFonts.set(qn(f'w:{tag}'), _FONT)


def _para(p, text, *, size=9, bold=False, align=None, color=None, space_after=4):
    p.alignment = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(str(text if text is not None else ''))
    _run_font(run, size=size, bold=bold, color=color)


def _add_para(doc, text, **kwargs):
    p = doc.add_paragraph()
    _para(p, text, **kwargs)
    return p


def _section_heading(doc, text):
    _add_para(doc, text, size=12, bold=True, color=_TITLE_COLOR, space_after=6)


def _cell_fill(cell, hex_color: str | None):
    if not hex_color:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def _cell(cell, text, *, size=8, bold=False, align=None, fill=None, color=None):
    cell.text = ''
    _para(cell.paragraphs[0], text, size=size, bold=bold, align=align, color=color, space_after=0)
    _cell_fill(cell, fill)


def _set_col_widths(table, widths_mm: list[float]):
    table.autofit = False
    if hasattr(table, 'allow_autofit'):
        table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_mm):
            row.cells[i].width = Mm(w)


def _tbl_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), 'CBD5E1')
        borders.append(el)
    tblPr.append(borders)


def _merge_row_groups(row, group_sizes: list[int]):
    """橫向合併儲存格（例：[2,2,2] 共 6 欄）"""
    idx = 0
    for span in group_sizes:
        if span > 1:
            row.cells[idx].merge(row.cells[idx + span - 1])
        idx += span


def _add_table(doc, rows: list[list], widths_mm: list[float], *,
               dark_header_rows: tuple[int, ...] = (),
               right_cols: tuple[int, ...] = (),
               label_cols: tuple[int, ...] = ()):
    if not rows:
        return None
    tbl = doc.add_table(rows=len(rows), cols=len(widths_mm))
    _set_col_widths(tbl, widths_mm)
    _tbl_borders(tbl)
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            align = 'right' if c in right_cols else None
            fill = None
            color = None
            bold = False
            if r in dark_header_rows:
                fill = _HDR_FILL
                color = _HDR_TEXT
                bold = True
            elif c in label_cols:
                fill = _LABEL_FILL
            _cell(tbl.rows[r].cells[c], text, bold=bold, align=align, fill=fill, color=color)
    return tbl


def _add_logo(doc):
    from sc_fac_pdf import _logo_draw_size, _logo_path
    path = _logo_path()
    if not path:
        return
    w, h = _logo_draw_size()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(path, width=Mm(w / rl_mm), height=Mm(h / rl_mm))


def _add_ip_period_table(doc, ip_items: list, ip_totals: dict):
    """糧期：摘要（三組 x2 欄）+ 明細表（同一表格 · 對照 QS Word）"""
    n_detail = len(ip_items)
    n_rows = 2 + (1 + n_detail if n_detail else 0)
    tbl = doc.add_table(rows=n_rows, cols=6)
    _set_col_widths(tbl, [18, 37, 30, 28, 32, 35])
    _tbl_borders(tbl)

    income = _money(ip_totals.get('total_income'))
    expend = _money(-abs(float(ip_totals.get('total_expenditure') or 0)))
    advance = _money(ip_totals.get('advance'))

    for label, val in zip(['總收入', '總支出', '墊支'], [income, expend, advance]):
        ci = ['總收入', '總支出', '墊支'].index(label) * 2
        _cell(tbl.rows[0].cells[ci], label, bold=True, fill=_HDR_FILL, color=_HDR_TEXT)
        _cell(tbl.rows[1].cells[ci], val)
    _merge_row_groups(tbl.rows[0], [2, 2, 2])
    _merge_row_groups(tbl.rows[1], [2, 2, 2])

    if not n_detail:
        return tbl

    hdr = ['期數', '申請日期', '申請金額', '累計%', '批款收入', '分包支出']
    for i, text in enumerate(hdr):
        align = 'right' if i >= 2 else None
        _cell(tbl.rows[2].cells[i], text, bold=True, fill=_HDR_FILL, color=_HDR_TEXT, align=align)

    for ri, it in enumerate(ip_items):
        row = tbl.rows[3 + ri]
        vals = [
            _plain(it.get('ip_no'), 8),
            _plain(it.get('applied_date'), 12),
            _money(it.get('application_amount')),
            _pct(it.get('application_pct')),
            _money(it.get('certified_income')),
            _money(it.get('subcon_paid')),
        ]
        for ci, text in enumerate(vals):
            align = 'right' if ci >= 2 else None
            _cell(row.cells[ci], text, align=align)
    return tbl


def generate_boss_qs_report_docx(summary: dict, sc_list: list | None = None,
                                 company_name: str = '', payment_count: int = 0) -> bytes:
    """生成 A4 QS 地盤財務匯報 Word（版面對照 QS 2026-08-14 交回版）"""
    sc_list = sc_list or []
    project = summary.get('project') or {}
    calc = summary.get('contract_calc') or {}
    ip = summary.get('ip_period') or {}
    ip_items = ip.get('items') or []
    ip_totals = ip.get('totals') or {}
    report_date = datetime.now().strftime('%Y-%m-%d %H:%M')
    code = project.get('project_code') or 'PROJECT'
    en, zh = _project_titles(project)

    contract_a = float(calc.get('main_contract_amount') or 0)
    total_paid = float(summary.get('total_paid') or 0)
    total_rem = float(summary.get('total_remainder') or 0)
    pay_progress = (total_paid / contract_a * 100) if contract_a else 0

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = _FONT
    style.font.size = Pt(9)
    try:
        style.element.rPr.rFonts.set(qn('w:eastAsia'), _FONT)
    except Exception:
        pass

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(15)
    sec.bottom_margin = Mm(16)
    sec.left_margin = Mm(15)
    sec.right_margin = Mm(15)

    _add_logo(doc)
    _add_para(doc, 'QS 地盤財務匯報', size=16, bold=True, align='center', space_after=2, color=_TITLE_COLOR)
    _add_para(doc, 'Quantity Surveying · Site Financial Summary', size=10, align='center',
              color=_MUTED_COLOR, space_after=8)

    _add_table(doc, [
        ['項目代碼', code, '報告日期', report_date],
        ['項目名稱（中）', _plain(zh or '—', 80), '客戶', _plain(project.get('client'))],
        ['項目名稱（英）', _plain(en or '—', 80), '主承建商', _plain(project.get('main_contractor'), 40)],
        ['工期', _plain(project.get('site_period_text') or ip.get('site_period_text'), 50),
         '狀態', STATUS_LABELS.get(project.get('status'), project.get('status') or '—')],
    ], [28, 62, 28, 52], label_cols=(0, 2))

    _section_heading(doc, '一、重點摘要')
    _add_table(doc, [
        ['承建金額 (A)', '預計利潤 (E)', '預計利潤率'],
        [_money(calc.get('main_contract_amount')), _money(calc.get('profit_e')), _pct(calc.get('profit_rate'))],
        ['累計已付', '未付餘額', '付款進度 / 墊支'],
        [_money(total_paid), _money(total_rem), f'{_pct(pay_progress)} / {_money(ip_totals.get("advance"))}'],
    ], [56, 56, 56], dark_header_rows=(0,))
    _add_para(doc,
              f'判項/支出 {len(sc_list)} 項 · 付款登記 {payment_count} 筆 · 糧期 {len(ip_items)} 期',
              size=8, color=_MUTED_COLOR, space_after=8)

    _section_heading(doc, '二、關注事項')
    for line in _attention_items(summary, sc_list):
        _add_para(doc, f'• {line}', size=9, space_after=2)

    _section_heading(doc, '三、合約金額結算 (A–E)')
    _add_table(doc, [
        ['項目', '金額 (HK$)'],
        ['(A) 承建金額', _money(calc.get('main_contract_amount'))],
        ['(B) 分判及代支小計', _money(calc.get('sub_total_b'))],
        ['(C) 除外合約收費項目', _money(calc.get('excluded_c'))],
        ['財務會作調撥（人工分攤）', _money(calc.get('labour_allocation'))],
        ['(D) = (B)+(C)+調撥', _money(calc.get('total_d'))],
        ['(E) = (A)−(D) 預計利潤', _money(calc.get('profit_e'))],
        ['預計利潤率', _pct(calc.get('profit_rate'))],
    ], [110, 58], dark_header_rows=(0,), right_cols=(1,))

    _section_heading(doc, '四、費用類別概覽')
    _add_table(doc, _category_rows(sc_list), [35, 45, 45, 43], dark_header_rows=(0,), right_cols=(1, 2, 3))

    _section_heading(doc, '五、地盤糧期狀況')
    if ip_items or ip_totals:
        _add_ip_period_table(doc, ip_items, ip_totals)
    else:
        _add_para(doc, '暫無糧期記錄', size=9, space_after=6)

    _section_heading(doc, '六、分判及支出明細')
    if sc_list:
        sc_rows = [['判項', '類別', '公司', '判項金額', '已付', '未付', '進度']]
        sum_ca = sum_paid = 0.0
        for sc in sorted(sc_list, key=lambda x: (x.get('sc_no') or '')):
            ca = float(sc.get('contract_amount') or 0)
            paid = float(sc.get('total_paid') or 0)
            rem = ca - paid
            pct = (paid / ca * 100) if ca else 0
            sum_ca += ca
            sum_paid += paid
            sc_rows.append([
                _plain(sc.get('sc_no'), 10),
                _sc_category(sc.get('sc_no')),
                _plain(_company_primary(sc), 28),
                _money(ca), _money(paid), _money(rem), _pct(pct),
            ])
        sc_rows.append(['合計', '', '', _money(sum_ca), _money(sum_paid), _money(sum_ca - sum_paid), ''])
        _add_table(doc, sc_rows, [15, 16, 27, 32, 32, 32, 14], dark_header_rows=(0,), right_cols=(3, 4, 5, 6))
    else:
        _add_para(doc, '暫無判項/支出項', size=9, space_after=6)

    _add_para(doc,
              f'本報告由 QS 管理系統自動生成 · {report_date} · 僅供內部管理參考',
              size=8, color=_MUTED_COLOR, space_after=0)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
