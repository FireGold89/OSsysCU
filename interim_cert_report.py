"""承判商工程中期糧款計算書 — 對齊 QS Excel 範本（駿昇第十期）輸出（PDF / XLSX / Word）"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, Side
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import Image as RLImage, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from qs_report_pdf import _esc, _plain, ensure_pdf_font, ensure_pdf_font_bold
from sc_vo_templates import (
    cert_label_for_record,
    get_cert_standard_lines,
    normalize_standard_amount,
)

MARGIN = 12 * mm
COMPANY = 'MEPORK ENGINEERING SERVICES LIMITED'
# 對齊使用者 Word 改版（A4 內容寬 186mm；右兩欄對齊備註）
META_W = [45 * mm, 77.5 * mm, 32 * mm, 31.5 * mm]
AMT_W = [45 * mm, 25 * mm, 27.5 * mm, 25 * mm, 63.5 * mm]
SIG_W = AMT_W
SUM_LABELS = ('總計', 'Sub-total (A+B)')
META_W_MM = [45, 77.5, 32, 31.5]
AMT_W_MM = [45, 25, 27.5, 25, 63.5]
HDR_FILL = 'E2E8F0'
HDR_TEXT = '1E293B'


def _f(val, default=0.0) -> float:
    try:
        return float(val)
    except (TypeError, ValueError):
        return default


def _money_plain(val) -> str:
    """表頭承包價：正數 1,234.56；零 0.00；負數 (1,234.56)"""
    if val is None or val == '':
        return '-'
    try:
        n = float(val)
    except (TypeError, ValueError):
        return '-'
    if n == 0:
        return '0.00'
    if n < 0:
        return f'({abs(n):,.2f})'
    return f'{n:,.2f}'


def _money_amt(val) -> str:
    """金額三欄：有數靠右顯示；0／空白為 -；負數 (1,234.56)"""
    if val is None or val == '':
        return '-'
    try:
        n = float(val)
    except (TypeError, ValueError):
        return '-'
    if n == 0:
        return '-'
    if n < 0:
        return f'({abs(n):,.2f})'
    return f'{n:,.2f}'


def _is_sum_line(label: str) -> bool:
    return (label or '') in SUM_LABELS


def _logo_fitted():
    """沿用分判結算 logo 框（最闊 75.6mm、最高 12mm），回傳符合原圖比例嘅實際寬高。"""
    from sc_fac_pdf import LOGO_H, LOGO_W_MAX, _logo_path
    path = _logo_path()
    if not path:
        return None, None, None
    try:
        from PIL import Image as PILImage
        with PILImage.open(path) as im:
            iw, ih = im.size
        if iw <= 0 or ih <= 0:
            return path, LOGO_W_MAX, LOGO_H
        aspect = iw / float(ih)
        if LOGO_W_MAX / LOGO_H >= aspect:
            h = LOGO_H
            w = h * aspect
        else:
            w = LOGO_W_MAX
            h = w / aspect
        return path, w, h
    except Exception:
        return path, LOGO_W_MAX, LOGO_H


def _line(label, cum_cur, cum_prev, cur, remark='') -> dict:
    return {
        'label': label,
        'cum_current': round(cum_cur, 2),
        'cum_previous': round(cum_prev, 2),
        'current_provisional': round(cur, 2),
        'remark': remark or '',
    }


def _retention_cumulative(sub_cum: float, sc_total: float, retention_pct: float) -> float:
    """駿昇第十期：保固金按合約總承包價上限累計（負數）"""
    if sub_cum <= 0:
        return 0.0
    cap = round(sc_total * retention_pct, 2)
    return -cap


def _format_work_done_pct(a_cum: float, sc_contract: float) -> str:
    if sc_contract <= 0:
        return ''
    pct = round(a_cum / sc_contract * 100, 1)
    return f'{pct:g}% Work Done'


def _auto_remark_b(vo_items: list) -> str:
    return '、'.join(
        _default_vo_remark(v) for v in vo_items if _default_vo_remark(v)
    )[:80]


def _default_vo_remark(vo: dict) -> str:
    desc = (vo.get('description') or '').strip()
    if desc:
        return desc[:80]
    return (vo.get('ref_no') or '').strip()[:80]


def _default_deduction_remark(d: dict) -> str:
    desc = (d.get('description') or '').strip()
    if desc:
        return desc[:80]
    return (d.get('ref_no') or '').strip()[:80]


def _parse_vo_remarks_entry(entry) -> dict:
    if isinstance(entry, str):
        return {'text': entry.strip(), 'auto': False}
    if isinstance(entry, dict):
        return {
            'text': (entry.get('text') or '').strip(),
            'auto': bool(entry.get('auto', False)),
        }
    return {'text': '', 'auto': True}


def _stored_vo_remark(stored: dict, vo_id) -> dict:
    if not vo_id:
        return {'text': '', 'auto': True}
    key = str(vo_id)
    entry = stored.get(key)
    if entry is None and vo_id is not None:
        try:
            entry = stored.get(int(vo_id))
        except (TypeError, ValueError):
            pass
    return _parse_vo_remarks_entry(entry)


def _resolve_vo_remarks(cert: dict, vo_items: list) -> dict[str, dict]:
    """每條 VO 備註：auto=true 用 ref_no／description 預設。"""
    stored = cert.get('vo_remarks') or {}
    out: dict[str, dict] = {}
    for v in vo_items:
        vid = v.get('id')
        if vid is None:
            continue
        key = str(vid)
        entry = _stored_vo_remark(stored, vid)
        text = _default_vo_remark(v) if entry['auto'] else entry['text']
        out[key] = {'text': text, 'auto': entry['auto']}
    return out


def _merged_b_remark(vo_remarks: dict[str, dict], vo_items: list) -> str:
    parts = []
    for v in vo_items:
        key = str(v.get('id') or '')
        text = (vo_remarks.get(key) or {}).get('text') or _default_vo_remark(v)
        if text:
            parts.append(text)
    return '、'.join(parts)[:80]


def _build_b_lines(
    cert: dict,
    vo_items: list,
    b_prev: float,
    b_prov: float,
    b_cur: float,
    vo_remarks: dict[str, dict],
    line_remark_b: str,
) -> list[dict]:
    """Phase 2：有 VO 時展開 B1/B2…；否則單一 B 行。"""
    expand = cert.get('b_expand_vo', True)
    if expand and vo_items:
        lines = []
        for i, v in enumerate(vo_items, 1):
            amt = abs(_f(v.get('amount')))
            key = str(v.get('id') or '')
            remark = (vo_remarks.get(key) or {}).get('text') or _default_vo_remark(v)
            ref = (v.get('ref_no') or '').strip()
            label = f'B{i}. 後加 / 改工程'
            if ref:
                label = f'B{i}. 後加 / 改工程 ({ref})'
            if i == 1:
                cum_prev = b_prev
                cum_cur = round(b_prev + amt, 2)
            else:
                cum_prev = 0.0
                cum_cur = amt
            lines.append(_line(label, cum_cur, cum_prev, amt, remark))
        return lines
    return [_line('B. 後加 / 改工程', b_cur, b_prev, b_prov, line_remark_b)]


def _parse_line_remarks_entry(entry) -> dict:
    if isinstance(entry, str):
        return {'text': entry.strip(), 'auto': False}
    if isinstance(entry, dict):
        return {
            'text': (entry.get('text') or '').strip(),
            'auto': bool(entry.get('auto', False)),
        }
    return {'text': '', 'auto': True}


def _resolve_line_remarks(cert: dict, a_cur: float, sc_contract: float, vo_items: list) -> dict[str, str]:
    """A/B/C 備註：auto=true 用系統預設；手改後保留 text 且不再覆蓋。"""
    stored = cert.get('line_remarks') or {}
    if cert.get('line_a_remark') and 'A' not in stored:
        stored = dict(stored)
        stored['A'] = {'text': cert['line_a_remark'], 'auto': False}

    auto_text = {
        'A': _format_work_done_pct(a_cur, sc_contract),
        'B': _auto_remark_b(vo_items),
        'C': '',
    }
    out: dict[str, str] = {}
    for key in ('A', 'B', 'C'):
        entry = _parse_line_remarks_entry(stored.get(key))
        out[key] = auto_text[key] if entry['auto'] else entry['text']
    return out


def build_interim_cert_model(cert: dict) -> dict:
    """由表單／VO 登記資料組裝與 QS Excel 一致嘅行項目"""
    project = cert.get('project') or {}
    prev = cert.get('previous_state') or {}
    sc_contract = _f(cert.get('sc_contract_sum') or cert.get('contract_amount'))
    vo_total = _f(cert.get('vo_amount'))
    sc_total = sc_contract + vo_total
    retention_pct = _f(cert.get('retention_pct'), 0.05)
    net = _f(cert.get('paid_amount'))

    a_prev = _f(cert.get('previous_a_cum', prev.get('a_cum')))
    b_prev = _f(cert.get('previous_b_cum', prev.get('b_cum')))
    ret_prev = _f(cert.get('previous_ret_cum', prev.get('ret_cum')))

    vo_items = cert.get('vo_items') or []
    b_prov = sum(_f(v.get('amount')) for v in vo_items)
    b_cur = b_prev + b_prov

    c_cur = c_prev = c_prov = 0.0

    contra_lines = []
    ded_prov_sum = 0.0
    for d in cert.get('deductions') or []:
        amt = _f(d.get('amount'))
        if amt >= 0:
            amt = -abs(amt)
        ded_prov_sum += amt
        label = cert_label_for_record(d) or f"減: {(d.get('description') or '扣款').strip()}"
        contra_lines.append(_line(label, amt, 0.0, amt, _default_deduction_remark(d)))

    standard_inputs = cert.get('standard_lines') or {}
    selected_std = cert.get('selected_standard_codes') or list(standard_inputs.keys())
    prev_std = (prev.get('standard_cums') or {}) if isinstance(prev, dict) else {}
    standard_rows = []
    std_prov_sum = 0.0
    for tpl in get_cert_standard_lines():
        code = tpl['code']
        if code not in selected_std:
            continue
        label = tpl['cert_label']
        raw = standard_inputs.get(code)
        prov = normalize_standard_amount(code, _f(raw)) if raw is not None and raw != '' else 0.0
        prev_cum = _f(prev_std.get(code))
        cur_cum = prev_cum + prov
        standard_rows.append(_line(label, cur_cum, prev_cum, prov))
        std_prov_sum += prov

    sub_prev = a_prev + b_prev + c_prev
    ret_prov_est = _retention_cumulative(sub_prev + b_prov, sc_total, retention_pct) - ret_prev

    a_prov = cert.get('a_current_provisional')
    if a_prov is None:
        a_prov = round(net - b_prov - ret_prov_est - ded_prov_sum - std_prov_sum, 2)
        if a_prov < 0:
            a_prov = 0.0
            net = round(b_prov + ret_prov_est + ded_prov_sum + std_prov_sum, 2)
    else:
        a_prov = _f(a_prov)

    a_cur = a_prev + a_prov
    sub_cur = a_cur + b_cur + c_cur
    sub_prov = a_prov + b_prov + c_prov

    ret_cur = _retention_cumulative(sub_cur, sc_total, retention_pct)
    ret_prov = round(ret_cur - ret_prev, 2)

    if cert.get('a_current_provisional') is None:
        a_prov = round(net - b_prov - ret_prov - ded_prov_sum - std_prov_sum, 2)
        if a_prov < 0:
            a_prov = 0.0
            net = round(b_prov + ret_prov + ded_prov_sum + std_prov_sum, 2)
        a_cur = a_prev + a_prov
        sub_cur = a_cur + b_cur + c_cur
        sub_prov = a_prov + b_prov + c_prov
        ret_cur = _retention_cumulative(sub_cur, sc_total, retention_pct)
        ret_prov = round(ret_cur - ret_prev, 2)

    total_cur = (
        sub_cur + ret_cur
        + sum(x['cum_current'] for x in contra_lines)
        + sum(x['cum_current'] for x in standard_rows)
    )
    total_prev = (
        sub_prev + ret_prev
        + sum(x['cum_previous'] for x in contra_lines)
        + sum(x['cum_previous'] for x in standard_rows)
    )
    total_prov = sub_prov + ret_prov + ded_prov_sum + std_prov_sum

    mp_sum = _f(cert.get('mp_contract_sum') or project.get('contract_amount'))
    vo_remarks = _resolve_vo_remarks(cert, vo_items)
    cert_lr = dict(cert)
    lr = dict(cert.get('line_remarks') or {})
    b_entry = _parse_line_remarks_entry(lr.get('B'))
    if vo_items and b_entry['auto']:
        lr['B'] = {'text': _merged_b_remark(vo_remarks, vo_items), 'auto': True}
        cert_lr['line_remarks'] = lr
    line_remarks = _resolve_line_remarks(cert_lr, a_cur, sc_contract, vo_items)
    b_lines = _build_b_lines(cert, vo_items, b_prev, b_prov, b_cur, vo_remarks, line_remarks['B'])

    lines = [
        _line('A. 原合約工程', a_cur, a_prev, a_prov, line_remarks['A']),
        *b_lines,
        _line('C. Material On Site', c_cur, c_prev, c_prov, line_remarks['C']),
        _line('Sub-total (A+B)', sub_cur, sub_prev, sub_prov),
        _line('減:保固金', ret_cur, ret_prev, ret_prov),
        *contra_lines,
        *standard_rows,
        _line('總計', total_cur, total_prev, total_prov),
    ]

    return {
        'company': COMPANY,
        'title': '承判商工程中期糧款計算書',
        'company_zh': cert.get('company_zh') or cert.get('company_name_zh') or '',
        'company_en': cert.get('company_en') or cert.get('company_name_en') or '',
        'invoice_no': cert.get('invoice_no') or '',
        'application_no': cert.get('application_no') or '',
        'invoice_date': cert.get('invoice_date') or '',
        'work_period': cert.get('work_period') or project.get('site_period_text') or '',
        'project_code': project.get('project_code') or cert.get('quotation_no') or '',
        'project_name': (
            project.get('project_name_en') or project.get('project_name_zh')
            or project.get('project_name') or cert.get('description') or ''
        ),
        'trade_label': cert.get('trade_label') or cert.get('description') or '',
        'mp_contract_sum': mp_sum,
        'sc_contract_sum': sc_contract,
        'vo_amount': vo_total,
        'sc_total_sum': sc_total,
        'table_header': [
            '累積至申請日期\n工程完成總額',
            '今期累計金額\nHK$',
            '上期累計金額\nHK$',
            '今期暫批金額\nHK$',
            '備註',
        ],
        'lines': lines,
        'prepared_by': cert.get('prepared_by') or project.get('person_in_charge') or '',
        'signature_date': cert.get('invoice_date') or datetime.now().strftime('%Y-%m-%d'),
        'vo_items': vo_items,
        'vo_remarks': vo_remarks,
        'b_expand_vo': cert.get('b_expand_vo', True),
        'retention_pct': retention_pct,
    }


def enrich_interim_cert_payload(cert: dict) -> dict:
    """從 DB 補充上期累計、VO 總額等（API 層呼叫）"""
    import database as db

    cert = dict(cert or {})
    project_id = cert.get('project_id') or (cert.get('project') or {}).get('id')
    sc_no = cert.get('sc_no')
    if project_id and sc_no:
        prev = db.get_previous_interim_state(project_id, sc_no, cert.get('exclude_payment_id'))
        cert.setdefault('previous_state', prev)
        cert.setdefault('previous_a_cum', prev['a_cum'])
        cert.setdefault('previous_b_cum', prev['b_cum'])
        cert.setdefault('previous_ret_cum', prev['ret_cum'])
        if not cert.get('previous_paid'):
            cert['previous_paid'] = prev['net_paid_total']
        vo_total = db.sum_sc_vo_amount(project_id, sc_no, 'vo')
        cert['vo_amount'] = vo_total

    vo_ids = cert.get('vo_ids') or []
    ded_ids = cert.get('deduction_ids') or []
    if vo_ids:
        rows = db.get_sc_vo_records_by_ids(vo_ids)
        cert['vo_items'] = [
            {'id': r['id'], 'ref_no': r['ref_no'], 'description': r['description'], 'amount': r['amount']}
            for r in rows if r.get('record_type') == 'vo'
        ]
    elif cert.get('vo_items'):
        pass
    else:
        cert.setdefault('vo_items', [])

    if ded_ids and not cert.get('deductions'):
        rows = db.get_sc_vo_records_by_ids(ded_ids)
        cert['deductions'] = [
            {
                'id': r['id'], 'ref_no': r['ref_no'], 'description': r['description'],
                'amount': r['amount'], 'line_code': r.get('line_code'),
            }
            for r in rows if r.get('record_type') == 'deduction'
        ]

    cert.setdefault('standard_lines', cert.get('standard_lines') or {})
    cert.setdefault('selected_standard_codes', cert.get('selected_standard_codes') or [])

    model = build_interim_cert_model(cert)
    cert['model'] = model
    return cert


def _pdf_styles(font_name: str) -> dict:
    bold_name = ensure_pdf_font_bold()
    return {
        'title': ParagraphStyle('ict', fontName=bold_name, fontSize=11, alignment=TA_CENTER, leading=14),
        'label': ParagraphStyle('icl', fontName=font_name, fontSize=8, leading=10),
        'label_r': ParagraphStyle('iclr', fontName=font_name, fontSize=8, leading=10, alignment=TA_RIGHT),
        'cell': ParagraphStyle('icc', fontName=font_name, fontSize=8, leading=10),
        'cell_r': ParagraphStyle('icr', fontName=font_name, fontSize=8, leading=10, alignment=TA_RIGHT),
        'cell_c': ParagraphStyle('icc2', fontName=font_name, fontSize=8, leading=10, alignment=TA_CENTER),
        'cell_b': ParagraphStyle('iccb', fontName=bold_name, fontSize=8, leading=10),
        'cell_rb': ParagraphStyle('icrb', fontName=bold_name, fontSize=8, leading=10, alignment=TA_RIGHT),
        'cell_cb': ParagraphStyle('icccb', fontName=bold_name, fontSize=8, leading=10, alignment=TA_CENTER),
        'hdr': ParagraphStyle('ich', fontName=bold_name, fontSize=8, leading=10, textColor=colors.HexColor('#' + HDR_TEXT)),
        'hdr_c': ParagraphStyle(
            'ichc', fontName=bold_name, fontSize=8, leading=10,
            alignment=TA_CENTER, textColor=colors.HexColor('#' + HDR_TEXT),
        ),
        'small': ParagraphStyle('ics', fontName=font_name, fontSize=8, textColor=colors.HexColor('#64748b')),
    }


def _p(text, style, *, bold=False) -> Paragraph:
    t = _esc(text).replace('\n', '<br/>')
    if bold:
        t = f'<b>{t}</b>'
    return Paragraph(t, style)


def _pdf_row_line(line: dict, styles: dict, *, bold=False, center_label=False):
    if bold:
        label_style = styles['cell_cb'] if center_label else styles['cell_b']
        num_style = styles['cell_rb']
        remark_style = styles['cell_b']
    else:
        label_style = styles['cell_c'] if center_label else styles['cell']
        num_style = styles['cell_r']
        remark_style = styles['cell']
    return [
        _p(line['label'], label_style),
        _p(_money_amt(line['cum_current']), num_style),
        _p(_money_amt(line['cum_previous']), num_style),
        _p(_money_amt(line['current_provisional']), num_style),
        _p(_plain(line.get('remark'), 80), remark_style),
    ]


def _pdf_logo_flowable():
    path, w, h = _logo_fitted()
    if not path or not w or not h:
        return None
    img = RLImage(path, width=w, height=h, mask='auto')
    img.hAlign = 'LEFT'
    return img


def generate_interim_cert_pdf(cert: dict) -> bytes:
    """生成對齊使用者 Word 改版嘅 A4 PDF"""
    enriched = enrich_interim_cert_payload(cert)
    model = enriched.get('model') or build_interim_cert_model(enriched)
    font_name = ensure_pdf_font()
    styles = _pdf_styles(font_name)
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=MARGIN, rightMargin=MARGIN,
        topMargin=14 * mm, bottomMargin=12 * mm,
        title='中期糧款計算書',
    )
    story = []

    logo = _pdf_logo_flowable()
    if logo:
        story.append(logo)
        story.append(Spacer(1, 10 * mm))
    story.append(_p(model['title'], styles['title']))
    story.append(Spacer(1, 4 * mm))

    _meta_pad = [
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]
    hdr = Table([
        [_p('承判人/公司名稱(中) :', styles['label']), _p(model['company_zh'], styles['cell']),
         _p('發票號碼 :', styles['label_r']), _p(model['invoice_no'], styles['cell_r'])],
        [_p('承判人/公司名稱(英) :', styles['label']), _p(model['company_en'], styles['cell']),
         _p('申請期數 :', styles['label_r']), _p(model['application_no'], styles['cell_r'])],
        [_p('工作期間 :', styles['label']), _p(model['work_period'], styles['cell']),
         _p('日期 :', styles['label_r']), _p(str(model['signature_date'])[:10], styles['cell_r'])],
    ], colWidths=META_W)
    hdr.setStyle(TableStyle([
        ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
        ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.HexColor('#e2e8f0')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f8fafc')),
        ('BACKGROUND', (2, 0), (2, -1), colors.HexColor('#f8fafc')),
        *_meta_pad,
    ]))
    story.append(hdr)
    story.append(Spacer(1, 3 * mm))

    proj = Table([
        [_p('工程編號:', styles['label']), _p(model['project_code'], styles['cell']),
         _p('合約承包價 :', styles['label_r']), _p(_money_plain(model['sc_contract_sum']), styles['cell_r'])],
        [_p('工程名稱 :', styles['label']), _p(_plain(model['project_name'], 90), styles['cell']),
         _p('後加/更改承包價 :', styles['label_r']), _p(_money_plain(model['vo_amount']), styles['cell_r'])],
        [_p('工程性質/工種 :', styles['label']), _p(_plain(model['trade_label'], 60), styles['cell']),
         _p('合約總承包價 :', styles['label_r']), _p(_money_plain(model['sc_total_sum']), styles['cell_r'])],
    ], colWidths=META_W)
    proj.setStyle(TableStyle([
        ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
        ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.HexColor('#e2e8f0')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f8fafc')),
        ('BACKGROUND', (2, 0), (2, -1), colors.HexColor('#f8fafc')),
        *_meta_pad,
    ]))
    story.append(proj)
    story.append(Spacer(1, 4 * mm))

    th = model['table_header']
    rows = [[
        _p(th[0], styles['hdr']),
        _p(th[1], styles['hdr_c']),
        _p(th[2], styles['hdr_c']),
        _p(th[3], styles['hdr_c']),
        _p(th[4], styles['hdr_c']),
    ]]
    sum_rows = []
    for i, line in enumerate(model['lines']):
        emph = _is_sum_line(line['label'])
        if emph:
            sum_rows.append(i + 1)
        rows.append(_pdf_row_line(line, styles, bold=emph, center_label=emph))

    tbl = Table(rows, colWidths=AMT_W, repeatRows=1)
    amt_style = [
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#' + HDR_FILL)),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#' + HDR_TEXT)),
        ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#94a3b8')),
        ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.HexColor('#cbd5e1')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        *_meta_pad,
    ]
    for r in sum_rows:
        amt_style.append(('BACKGROUND', (0, r), (-1, r), colors.HexColor('#f1f5f9')))
    tbl.setStyle(TableStyle(amt_style))
    story.append(tbl)
    story.append(Spacer(1, 5 * mm))

    sig = Table([
        [_p('内部使用', styles['cell_c']), _p('Prepared by', styles['cell_c']),
         _p('QS', styles['cell_c']), _p('PM', styles['cell_c']), _p('General Manager', styles['cell_c'])],
        [_p('簽署', styles['cell_c']), _p(model['prepared_by'], styles['cell_c']), _p('', styles['cell']),
         _p('', styles['cell']), _p('', styles['cell'])],
        [_p('日期', styles['cell_c']), _p(str(model['signature_date'])[:10], styles['cell_c']),
         _p('', styles['cell']), _p('', styles['cell']), _p('', styles['cell'])],
    ], colWidths=SIG_W, rowHeights=[8 * mm, 13 * mm, 8 * mm])
    sig.setStyle(TableStyle([
        ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
        ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.HexColor('#e2e8f0')),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f8fafc')),
        ('BACKGROUND', (0, 1), (0, -1), colors.HexColor('#f8fafc')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        *_meta_pad,
    ]))
    story.append(sig)

    doc.build(story)
    return buf.getvalue()


def _xlsx_border():
    thin = Side(style='thin', color='CCCCCC')
    return Border(left=thin, right=thin, top=thin, bottom=thin)


def _xlsx_num(cell, val):
    cell.value = val if val else None
    cell.number_format = '#,##0.00;($#,##0.00)'
    cell.alignment = Alignment(horizontal='right', vertical='top')


def generate_interim_cert_xlsx(cert: dict) -> bytes:
    """生成對齊駿昇第十期 Excel 範本嘅 xlsx bytes"""
    enriched = enrich_interim_cert_payload(cert)
    model = enriched.get('model') or build_interim_cert_model(enriched)
    wb = Workbook()
    ws = wb.active
    ws.title = f"IP Cert. {model['application_no']}"[:31]

    title_font = Font(bold=True, size=14)
    hdr_font = Font(bold=True, size=10)
    border = _xlsx_border()
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left = Alignment(horizontal='left', vertical='top', wrap_text=True)

    # 駿昇範本：工程編號在右上
    ws['E1'] = model['project_code'].split('_')[0].replace('MS_', 'Q').replace('_', '/')[:20] if model['project_code'] else ''
    ws.merge_cells('A3:E3')
    ws['A3'] = model['title']
    ws['A3'].font = title_font
    ws['A3'].alignment = center

    ws['A4'], ws['B4'] = '承判人/公司名稱(中) :', model['company_zh']
    ws['E4'], ws['F4'] = '發票號碼 :', model['invoice_no']
    ws['A5'], ws['B5'] = '承判人/公司名稱(英) :', model['company_en']
    ws['E5'], ws['F5'] = '申請期數 :', model['application_no']
    ws['A6'], ws['B6'] = '工作期間 :', model['work_period']
    ws['E6'], ws['F6'] = '日期 :', str(model['signature_date'])[:10]

    ws['A8'], ws['B8'] = '工程編號:', model['project_code']
    _xlsx_num(ws['F8'], model['sc_contract_sum'])
    ws['E8'] = '合約承包價 :'
    ws['A9'], ws['B9'] = '工程名稱 :', model['project_name']
    _xlsx_num(ws['F9'], model['vo_amount'])
    ws['E9'] = '後加/更改承包價 :'
    ws['A11'], ws['B11'] = '工程性質/工種 :', model['trade_label']
    _xlsx_num(ws['F11'], model['sc_total_sum'])
    ws['E11'] = '合約總承包價 :'

    headers = model['table_header']
    for col, h in enumerate(headers, 1):
        c = ws.cell(row=13, column=col, value=h.replace('\\n', '\n'))
        c.font = hdr_font
        c.alignment = center
        c.border = border

    row = 14
    for line in model['lines']:
        ws.cell(row=row, column=1, value=line['label']).alignment = left
        for col, key in enumerate(['cum_current', 'cum_previous', 'current_provisional'], 2):
            _xlsx_num(ws.cell(row=row, column=col), line.get(key))
            ws.cell(row=row, column=col).border = border
        ws.cell(row=row, column=5, value=line.get('remark') or '').alignment = left
        ws.cell(row=row, column=1).border = border
        ws.cell(row=row, column=5).border = border
        if line['label'] == '總計':
            for col in range(1, 6):
                ws.cell(row=row, column=col).font = Font(bold=True)
        row += 1

    sig_row = row + 1
    ws.cell(row=sig_row, column=1, value='内部使用')
    ws.cell(row=sig_row, column=2, value='Prepared by')
    ws.cell(row=sig_row, column=3, value='QS')
    ws.cell(row=sig_row, column=4, value='PM')
    ws.cell(row=sig_row, column=5, value='General Manager')
    ws.cell(row=sig_row + 1, column=1, value='簽署')
    ws.cell(row=sig_row + 1, column=2, value=model['prepared_by'])
    ws.cell(row=sig_row + 2, column=1, value='日期')
    ws.cell(row=sig_row + 2, column=2, value=str(model['signature_date'])[:10])

    ws.column_dimensions['A'].width = 32
    for col in 'BCDE':
        ws.column_dimensions[col].width = 18
    ws.column_dimensions['F'].width = 16

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


# ─── Word（供 QS 改版後交回調格式）──────────────────────────
_DOCX_FONT = 'Microsoft JhengHei'


def _docx_rgb(hex_color: str):
    from docx.shared import RGBColor
    h = hex_color.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _docx_run_font(run, *, size=9, bold=False, color=None):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
    run.font.name = _DOCX_FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = _docx_rgb(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), _DOCX_FONT)
    rFonts.set(qn('w:hAnsi'), _DOCX_FONT)
    rFonts.set(qn('w:eastAsia'), _DOCX_FONT)


def _docx_para(p, text, *, size=9, bold=False, align=None, color=None, space_after=0, space_before=0):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    p.alignment = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.08
    lines = str(text if text is not None else '').split('\n')
    for i, line in enumerate(lines):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        _docx_run_font(run, size=size, bold=bold, color=color)
    return p


def _docx_add_para(doc, text, **kwargs):
    p = doc.add_paragraph()
    return _docx_para(p, text, **kwargs)


def _docx_cell(cell, text, *, size=9, bold=False, align=None, color=None, fill=None, valign='top'):
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    cell.text = ''
    cell.vertical_alignment = {
        'center': WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        'bottom': WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
    }.get(valign, WD_CELL_VERTICAL_ALIGNMENT.TOP)
    _docx_para(cell.paragraphs[0], text, size=size, bold=bold, align=align, color=color)
    if fill:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill.lstrip('#'))
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)


def _docx_tbl_borders(table, outer='94A3B8', inner='CBD5E1'):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for edge, col, sz in (
        ('top', outer, '8'),
        ('left', outer, '8'),
        ('bottom', outer, '8'),
        ('right', outer, '8'),
        ('insideH', inner, '4'),
        ('insideV', inner, '4'),
    ):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), col)
        borders.append(el)
    tblPr.append(borders)


def _docx_set_widths(table, widths_mm):
    from docx.shared import Mm
    table.autofit = False
    if hasattr(table, 'allow_autofit'):
        table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_mm):
            row.cells[i].width = Mm(w)


def _docx_meta_table(doc, rows, widths_mm, *, right_cols=(), label_fill='F8FAFC'):
    tbl = doc.add_table(rows=len(rows), cols=4)
    _docx_set_widths(tbl, widths_mm)
    _docx_tbl_borders(tbl, outer='CBD5E1', inner='E2E8F0')
    for r, cells in enumerate(rows):
        for c, text in enumerate(cells):
            align = 'right' if c in right_cols else None
            fill = label_fill if c in (0, 2) else None
            _docx_cell(tbl.rows[r].cells[c], text, size=8, align=align, fill=fill)
    return tbl


def generate_interim_cert_docx(cert: dict) -> bytes:
    """生成對齊使用者 Word 改版嘅 A4 計算書。"""
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt

    enriched = enrich_interim_cert_payload(cert)
    model = enriched.get('model') or build_interim_cert_model(enriched)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = _DOCX_FONT
    style.font.size = Pt(8)
    try:
        style.element.rPr.rFonts.set(qn('w:eastAsia'), _DOCX_FONT)
    except Exception:
        pass

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(14)
    sec.bottom_margin = Mm(12)
    sec.left_margin = Mm(12)
    sec.right_margin = Mm(12)

    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = Pt(28)
    path, w, h = _logo_fitted()
    if path and w and h:
        p0.add_run().add_picture(path, width=Pt(w), height=Pt(h))
    _docx_add_para(doc, model['title'], size=11, bold=True, align='center', space_after=8)

    _docx_meta_table(doc, [
        ('承判人/公司名稱(中) :', model['company_zh'], '發票號碼 :', model['invoice_no']),
        ('承判人/公司名稱(英) :', model['company_en'], '申請期數 :', model['application_no']),
        ('工作期間 :', model['work_period'], '日期 :', str(model['signature_date'])[:10]),
    ], META_W_MM, right_cols=(2, 3))
    _docx_add_para(doc, '', size=6, space_after=4)

    _docx_meta_table(doc, [
        ('工程編號:', model['project_code'], '合約承包價 :', _money_plain(model['sc_contract_sum'])),
        ('工程名稱 :', model['project_name'], '後加/更改承包價 :', _money_plain(model['vo_amount'])),
        ('工程性質/工種 :', model['trade_label'], '合約總承包價 :', _money_plain(model['sc_total_sum'])),
    ], META_W_MM, right_cols=(2, 3))
    _docx_add_para(doc, '', size=6, space_after=6)

    th = model['table_header']
    lines = model['lines']
    tbl = doc.add_table(rows=1 + len(lines), cols=5)
    _docx_set_widths(tbl, AMT_W_MM)
    _docx_tbl_borders(tbl)
    for i, h in enumerate(th):
        _docx_cell(
            tbl.rows[0].cells[i], h.replace('\\n', '\n'),
            size=8, bold=True, align=None if i == 0 else 'center',
            color=HDR_TEXT, fill=HDR_FILL, valign='center',
        )
    for r, line in enumerate(lines, 1):
        emph = _is_sum_line(line['label'])
        fill = 'F1F5F9' if emph else None
        _docx_cell(
            tbl.rows[r].cells[0], line['label'], size=8, bold=emph,
            align='center' if emph else None, fill=fill,
        )
        _docx_cell(tbl.rows[r].cells[1], _money_amt(line['cum_current']), size=8, bold=emph, align='right', fill=fill)
        _docx_cell(tbl.rows[r].cells[2], _money_amt(line['cum_previous']), size=8, bold=emph, align='right', fill=fill)
        _docx_cell(tbl.rows[r].cells[3], _money_amt(line['current_provisional']), size=8, bold=emph, align='right', fill=fill)
        _docx_cell(tbl.rows[r].cells[4], line.get('remark') or '', size=8, fill=fill)

    _docx_add_para(doc, '', size=6, space_after=8)

    sig = doc.add_table(rows=3, cols=5)
    _docx_set_widths(sig, AMT_W_MM)
    _docx_tbl_borders(sig, outer='CBD5E1', inner='E2E8F0')
    headers = ['内部使用', 'Prepared by', 'QS', 'PM', 'General Manager']
    for i, h in enumerate(headers):
        _docx_cell(sig.rows[0].cells[i], h, size=8, bold=(i == 0), align='center', fill='F8FAFC')
    _docx_cell(sig.rows[1].cells[0], '簽署', size=8, align='center', fill='F8FAFC')
    _docx_cell(sig.rows[1].cells[1], model['prepared_by'], size=8, align='center')
    for i in range(2, 5):
        _docx_cell(sig.rows[1].cells[i], '', size=8, align='center')
    sig.rows[1].height = Mm(13)
    sig.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    _docx_cell(sig.rows[2].cells[0], '日期', size=8, align='center', fill='F8FAFC')
    _docx_cell(sig.rows[2].cells[1], str(model['signature_date'])[:10], size=8, align='center')
    for i in range(2, 5):
        _docx_cell(sig.rows[2].cells[i], '', size=8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
